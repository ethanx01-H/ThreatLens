"""
Professional DOCX Report Generator
Generates threat intelligence investigation reports matching the
IC Cybersecurity Incident Report template style.
"""

import os
from datetime import datetime
from typing import Dict, List, Any, Optional

import os as _os

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Logo paths (relative to this file or exe bundle)
_BUNDLE_DIR = _os.path.dirname(_os.path.abspath(__file__))
_LOGO_512 = _os.path.join(_BUNDLE_DIR, "threatlens.png")
_WATERMARK = _os.path.join(_BUNDLE_DIR, "threatlens.png")


# Country code -> full name mapping (common codes in threat intel)
_COUNTRY_NAMES = {
    "AF": "Afghanistan", "AL": "Albania", "DZ": "Algeria", "AD": "Andorra",
    "AO": "Angola", "AR": "Argentina", "AM": "Armenia", "AU": "Australia",
    "AT": "Austria", "AZ": "Azerbaijan", "BS": "Bahamas", "BH": "Bahrain",
    "BD": "Bangladesh", "BB": "Barbados", "BY": "Belarus", "BE": "Belgium",
    "BZ": "Belize", "BJ": "Benin", "BT": "Bhutan", "BO": "Bolivia",
    "BA": "Bosnia and Herzegovina", "BW": "Botswana", "BR": "Brazil",
    "BN": "Brunei", "BG": "Bulgaria", "BF": "Burkina Faso", "BI": "Burundi",
    "KH": "Cambodia", "CM": "Cameroon", "CA": "Canada", "CF": "Central African Republic",
    "TD": "Chad", "CL": "Chile", "CN": "China", "CO": "Colombia",
    "KM": "Comoros", "CG": "Congo", "CD": "Congo (DR)", "CR": "Costa Rica",
    "CI": "Côte d'Ivoire", "HR": "Croatia", "CU": "Cuba", "CY": "Cyprus",
    "CZ": "Czech Republic", "DK": "Denmark", "DJ": "Djibouti", "DO": "Dominican Republic",
    "EC": "Ecuador", "EG": "Egypt", "SV": "El Salvador", "GQ": "Equatorial Guinea",
    "ER": "Eritrea", "EE": "Estonia", "ET": "Ethiopia", "FJ": "Fiji",
    "FI": "Finland", "FR": "France", "GA": "Gabon", "GM": "Gambia",
    "GE": "Georgia", "DE": "Germany", "GH": "Ghana", "GR": "Greece",
    "GT": "Guatemala", "GN": "Guinea", "GY": "Guyana", "HT": "Haiti",
    "HN": "Honduras", "HK": "Hong Kong", "HU": "Hungary", "IS": "Iceland",
    "IN": "India", "ID": "Indonesia", "IR": "Iran", "IQ": "Iraq",
    "IE": "Ireland", "IL": "Israel", "IT": "Italy", "JM": "Jamaica",
    "JP": "Japan", "JO": "Jordan", "KZ": "Kazakhstan", "KE": "Kenya",
    "KR": "South Korea", "KP": "North Korea", "KW": "Kuwait", "KG": "Kyrgyzstan",
    "LA": "Laos", "LV": "Latvia", "LB": "Lebanon", "LS": "Lesotho",
    "LR": "Liberia", "LY": "Libya", "LI": "Liechtenstein", "LT": "Lithuania",
    "LU": "Luxembourg", "MO": "Macau", "MK": "North Macedonia", "MG": "Madagascar",
    "MW": "Malawi", "MY": "Malaysia", "MV": "Maldives", "ML": "Mali",
    "MT": "Malta", "MR": "Mauritania", "MU": "Mauritius", "MX": "Mexico",
    "MD": "Moldova", "MC": "Monaco", "MN": "Mongolia", "ME": "Montenegro",
    "MA": "Morocco", "MZ": "Mozambique", "MM": "Myanmar", "NA": "Namibia",
    "NP": "Nepal", "NL": "Netherlands", "NZ": "New Zealand", "NI": "Nicaragua",
    "NE": "Niger", "NG": "Nigeria", "NO": "Norway", "OM": "Oman",
    "PK": "Pakistan", "PA": "Panama", "PG": "Papua New Guinea", "PY": "Paraguay",
    "PE": "Peru", "PH": "Philippines", "PL": "Poland", "PT": "Portugal",
    "QA": "Qatar", "RO": "Romania", "RU": "Russia", "RW": "Rwanda",
    "SA": "Saudi Arabia", "SN": "Senegal", "RS": "Serbia", "SG": "Singapore",
    "SK": "Slovakia", "SI": "Slovenia", "SO": "Somalia", "ZA": "South Africa",
    "SS": "South Sudan", "ES": "Spain", "LK": "Sri Lanka", "SD": "Sudan",
    "SR": "Suriname", "SE": "Sweden", "CH": "Switzerland", "SY": "Syria",
    "TW": "Taiwan", "TJ": "Tajikistan", "TZ": "Tanzania", "TH": "Thailand",
    "TL": "Timor-Leste", "TG": "Togo", "TT": "Trinidad and Tobago", "TN": "Tunisia",
    "TR": "Turkey", "TM": "Turkmenistan", "UG": "Uganda", "UA": "Ukraine",
    "AE": "United Arab Emirates", "GB": "United Kingdom", "US": "United States",
    "UY": "Uruguay", "UZ": "Uzbekistan", "VE": "Venezuela", "VN": "Vietnam",
    "YE": "Yemen", "ZM": "Zambia", "ZW": "Zimbabwe",
    # Common non-standard codes
    "XK": "Kosovo", "EU": "Europe", "AP": "Asia Pacific",
}


def _country_name(code: str) -> str:
    """Convert 2-letter country code to full name. Returns code if unknown."""
    if not code:
        return "N/A"
    upper = code.strip().upper()
    return _COUNTRY_NAMES.get(upper, upper)


# ═══════════════════════════════════════════════════════════════════
# Design Constants
# ═══════════════════════════════════════════════════════════════════

FONT_NAME = "Century Gothic"
TITLE_COLOR = RGBColor(0x00, 0x10, 0x33)       # #001033 dark navy
HEADER_BG = "EB6E19"                            # orange section header
LABEL_COLOR = RGBColor(0x3B, 0x38, 0x38)       # #3B3838 field labels
ANSWER_BG = "F2F2F2"                            # light gray answer cells
MARGIN = Cm(1.27)                                # 0.5 inch

SEV_COLORS = {
    "CRITICAL": ("E74C3C", True),
    "HIGH":     ("E67E22", True),
    "MEDIUM":   ("F1C40F", False),
    "LOW":      ("27AE60", True),
    "INFO":     ("3498DB", True),
}


# ═══════════════════════════════════════════════════════════════════
# Styling Helpers
# ═══════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, color_hex):
    """Apply background color to a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right, each a dict with sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        el = OxmlElement(f'w:{edge}')
        for attr_name, attr_val in attrs.items():
            el.set(qn(f'w:{attr_name}'), str(attr_val))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _merge_row(table, row_idx, ncols):
    """Merge all cells in a row into one cell spanning ncols."""
    row = table.rows[row_idx]
    first = row.cells[0]
    last = row.cells[ncols - 1]
    first.merge(last)
    return first


def _ct(cell, text, bold=False, size=Pt(10), color=None, align=None, font_name=FONT_NAME, bg=None):
    """Write styled text to a table cell."""
    cell.text = str(text)
    if bg:
        _set_cell_shading(cell, bg)
    for p in cell.paragraphs:
        if align:
            p.alignment = align
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = size
            r.bold = bold
            r.font.name = font_name
            if color:
                r.font.color.rgb = color


def _add_styled_paragraph(doc, text, style=None, bold=False, size=Pt(10),
                          color=None, alignment=None, space_after=Pt(4),
                          font_name=FONT_NAME):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _make_section_table(doc, title, data_rows, ncols=2):
    """
    Create a section table with an orange merged header row and label/value rows.
    data_rows: list of (label, value) tuples.
    Returns the table.
    """
    nrows = 1 + len(data_rows)  # header + data rows
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Merged orange header row
    header_cell = _merge_row(table, 0, ncols)
    _set_cell_shading(header_cell, HEADER_BG)
    _ct(header_cell, title, bold=True, size=Pt(12),
        color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.LEFT)

    # Data rows: label in col 0, value in col 1 with gray bg
    for i, (label, value) in enumerate(data_rows):
        row = table.rows[i + 1]
        _ct(row.cells[0], label, bold=True, size=Pt(10), color=LABEL_COLOR)
        _ct(row.cells[1], str(value), size=Pt(10), bg=ANSWER_BG)

    return table


def _make_header_table(doc, title, headers, data, ncols=None):
    """
    Create a section table with orange merged header + column headers + data rows.
    headers: list of column header strings.
    data: list of tuples, each a row.
    """
    ncols = ncols or len(headers)
    nrows = 1 + 1 + len(data)  # section header + col headers + data
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Merged orange section header
    header_cell = _merge_row(table, 0, ncols)
    _set_cell_shading(header_cell, HEADER_BG)
    _ct(header_cell, title, bold=True, size=Pt(12),
        color=RGBColor(255, 255, 255), align=WD_ALIGN_PARAGRAPH.LEFT)

    # Column headers
    for j, h in enumerate(headers):
        _ct(table.rows[1].cells[j], h, bold=True, size=Pt(9),
            color=LABEL_COLOR)

    # Data rows
    for i, row_data in enumerate(data):
        row = table.rows[i + 2]
        for j, val in enumerate(row_data):
            _ct(row.cells[j], str(val), size=Pt(9), bg=ANSWER_BG)

    return table


def _add_watermark(doc):
    """Add a semi-transparent logo watermark behind page content."""
    if not _os.path.exists(_WATERMARK):
        return
    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(_WATERMARK, width=Cm(6))


# ═══════════════════════════════════════════════════════════════════
# Report Generator
# ═══════════════════════════════════════════════════════════════════

def generate_report(
    target: str,
    target_type: str,  # "ip" or "domain"
    risk_assessment: dict,
    ipinfo: dict = None,
    otx: dict = None,
    abuseipdb: dict = None,
    vt: dict = None,
    shodan: dict = None,
    threatfox: dict = None,
    urlhaus: dict = None,
    recon: dict = None,
    output_path: str = None,
    analyst: str = "Threat Intel Analyst",
    classification: str = "CONFIDENTIAL",
) -> str:
    """Generate a professional DOCX threat intelligence report."""

    if not HAS_DOCX:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Page margins: 0.5 inch
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    report_id = f"TI-{datetime.now().strftime('%Y%m%d')}-{target.replace('.', '-').replace(':', '-')}"
    risk_cls = risk_assessment.get("classification", "LOW")
    risk_score = risk_assessment.get("score", 0)

    # ─── Title ──────────────────────────────────────────────────
    _add_styled_paragraph(doc, "ThreatLens Report",
                          bold=True, size=Pt(24), color=TITLE_COLOR,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(4))

    # Confidential header
    _add_styled_paragraph(doc, "Confidential \u2014 For Internal Use Only",
                          bold=True, size=Pt(11), color=LABEL_COLOR,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(12))

    # ─── Report Information ─────────────────────────────────────
    _make_section_table(doc, "Report Information", [
        ("Report ID", report_id),
        ("Date", timestamp),
        ("Prepared By", analyst),
        ("Classification", classification),
        ("Report Type", "IP Address Analysis" if target_type == "ip" else "Domain Analysis"),
    ])

    doc.add_paragraph()

    # ─── Investigation Summary ──────────────────────────────────
    _make_section_table(doc, "Investigation Summary", [
        ("Target", target),
        ("Target Type", target_type.upper()),
        ("Detection Method", "OSINT Multi-Source Aggregation"),
        ("Description", (
            f"Investigation of {target_type} indicator \"{target}\" across "
            f"AlienVault OTX, AbuseIPDB, VirusTotal, Shodan, ThreatFox, and URLhaus."
        )),
    ])

    doc.add_paragraph()

    # ─── IPInfo / Geolocation ─────────────────────────────────
    if ipinfo and not ipinfo.get("error"):
        geo_data = [
            ("IP Address", ipinfo.get("ip", target)),
            ("Hostname", ipinfo.get("hostname", "N/A") or "N/A"),
            ("Country", _country_name(ipinfo.get("country", "N/A"))),
            ("City / Region", f"{ipinfo.get('city', '')}, {ipinfo.get('region', '')}".strip(", ")),
            ("Coordinates", ipinfo.get("loc", "N/A")),
            ("ASN", ipinfo.get("asn", "N/A")),
            ("ISP / Organization", ipinfo.get("isp", "N/A")),
            ("Timezone", ipinfo.get("timezone", "N/A")),
        ]
        if ipinfo.get("is_cloud"):
            geo_data.append(("Cloud Provider", ipinfo.get("cloud_provider", "")))
        _make_section_table(doc, "IPInfo / Geolocation", geo_data)
        doc.add_paragraph()

    # ─── Risk Assessment ────────────────────────────────────────
    signals = risk_assessment.get("signals", [])
    top_signals = signals[:3] if signals else []
    signal_text = "; ".join(s.get("signal", "") for s in top_signals) if top_signals else "None"

    otx_pulses = str(otx.get("pulse_count", 0) if otx and not otx.get("error") else "N/A")
    vt_mal = vt.get("malicious", 0) if vt and not vt.get("error") else 0
    vt_total = (vt_mal + vt.get("harmless", 0) + vt.get("undetected", 0) + vt.get("suspicious", 0)) if vt and not vt.get("error") else 0
    vt_ratio = f"{vt_mal}/{vt_total}" if vt and not vt.get("error") else "N/A"
    abuse_score = f"{abuseipdb.get('abuse_confidence_score', 0)}%" if abuseipdb and not abuseipdb.get("error") else "N/A"

    risk_table = _make_section_table(doc, "Risk Assessment", [
        ("Risk Score", f"{risk_score}/100"),
        ("Classification", risk_cls),
        ("Signals Detected", str(risk_assessment.get("signal_count", 0))),
        ("OTX Pulses", otx_pulses),
        ("VT Detections", vt_ratio),
        ("Abuse Score", abuse_score),
        ("Top Signals", signal_text),
    ])

    # Color-code the classification value cell
    if risk_cls in SEV_COLORS:
        hex_color, white_text = SEV_COLORS[risk_cls]
        _set_cell_shading(risk_table.rows[2].cells[1], hex_color)
        tc = RGBColor(255, 255, 255) if white_text else RGBColor(0, 0, 0)
        _ct(risk_table.rows[2].cells[1], risk_cls, bold=True, size=Pt(10), color=tc)

    doc.add_paragraph()

    # ─── Signal Analysis ────────────────────────────────────────
    if signals:
        sig_data = []
        for i, sig in enumerate(signals):
            sig_data.append((
                str(i + 1),
                sig.get("source", ""),
                sig.get("signal", ""),
                sig.get("severity", "INFO"),
            ))
        sig_table = _make_header_table(doc, "Signal Analysis",
                                        ["#", "Source", "Signal", "Severity"],
                                        sig_data, ncols=4)
        # Color severity cells
        for i, sig in enumerate(signals):
            sev = sig.get("severity", "INFO")
            hex_color, white_text = SEV_COLORS.get(sev, ("95A5A6", True))
            row_idx = i + 2  # skip section header + column header
            _set_cell_shading(sig_table.rows[row_idx].cells[3], hex_color)
            tc = RGBColor(255, 255, 255) if white_text else RGBColor(0, 0, 0)
            _ct(sig_table.rows[row_idx].cells[3], sev, bold=True, size=Pt(9), color=tc)
        doc.add_paragraph()

    # ─── Threat Intelligence Findings ───────────────────────────
    # Source summary table
    ti_rows = []

    if otx and not otx.get("error"):
        pc = otx.get("pulse_count", 0)
        ti_rows.append(("AlienVault OTX", "MATCH" if pc > 0 else "NO MATCH",
                        f"{pc} pulse(s)", f"Malware: {otx.get('malware_count', 0)}, URLs: {otx.get('url_count', 0)}"))
    else:
        ti_rows.append(("AlienVault OTX", otx.get("status", "NOT RUN") if otx else "NOT RUN", "-", ""))

    if abuseipdb and not abuseipdb.get("error"):
        acs = abuseipdb.get("abuse_confidence_score", 0)
        ti_rows.append(("AbuseIPDB", "MATCH" if acs > 0 else "NO MATCH",
                        f"{acs}%", f"Reports: {abuseipdb.get('total_reports', 0)}"))
    else:
        ti_rows.append(("AbuseIPDB", abuseipdb.get("status", "NOT RUN") if abuseipdb else "NOT RUN", "-", ""))

    if vt and not vt.get("error"):
        mal = vt.get("malicious", 0)
        sus = vt.get("suspicious", 0)
        total_vt = mal + sus + vt.get("harmless", 0) + vt.get("undetected", 0)
        status_vt = "MATCH" if mal > 0 else ("SUSPICIOUS" if sus > 0 else "NO MATCH")
        ti_rows.append(("VirusTotal", status_vt,
                        f"{mal}/{total_vt} detections", f"Undetected: {vt.get('undetected', 0)}"))
    else:
        ti_rows.append(("VirusTotal", vt.get("status", "NOT RUN") if vt else "NOT RUN", "-", ""))

    if shodan and not shodan.get("error"):
        ports = shodan.get("ports", [])
        vulns = shodan.get("vulns", [])
        ti_rows.append(("Shodan", "MATCH" if ports else "NO MATCH",
                        f"{len(ports)} port(s)", f"CVEs: {len(vulns)}"))
    else:
        ti_rows.append(("Shodan", shodan.get("status", "NOT RUN") if shodan else "NOT RUN", "-", ""))

    if threatfox and not threatfox.get("error"):
        ioc_count = threatfox.get("ioc_count", 0)
        ti_rows.append(("ThreatFox", "MATCH" if ioc_count > 0 else "NO MATCH",
                        f"{ioc_count} IOC(s)", ""))
    else:
        ti_rows.append(("ThreatFox", threatfox.get("status", "NOT RUN") if threatfox else "NOT RUN", "-", ""))

    if urlhaus and not urlhaus.get("error"):
        is_listed = urlhaus.get("is_listed", False)
        ti_rows.append(("URLhaus", "MATCH" if is_listed else "NO MATCH",
                        f"URLs: {urlhaus.get('url_count', 0)}", f"Threat: {urlhaus.get('threat', 'N/A')}"))
    else:
        ti_rows.append(("URLhaus", urlhaus.get("status", "NOT RUN") if urlhaus else "NOT RUN", "-", ""))

    coverage = risk_assessment.get("coverage", {})
    if coverage:
        ti_rows.append(("Coverage", "-",
                        f"{coverage.get('checked', 0)}/{coverage.get('total', 0)}",
                        f"{coverage.get('percentage', 0)}% confidence"))

    ti_table = _make_header_table(doc, "Threat Intelligence Findings",
                                   ["Source", "Status", "Result", "Details"],
                                   ti_rows, ncols=4)

    # Color status cells
    for i, (src, status, result, details) in enumerate(ti_rows):
        row_idx = i + 2
        if status == "MATCH":
            _set_cell_shading(ti_table.rows[row_idx].cells[1], "E74C3C")
            _ct(ti_table.rows[row_idx].cells[1], status, bold=True, size=Pt(9),
                color=RGBColor(255, 255, 255))
        elif status == "NO MATCH":
            _set_cell_shading(ti_table.rows[row_idx].cells[1], "27AE60")
            _ct(ti_table.rows[row_idx].cells[1], status, bold=True, size=Pt(9),
                color=RGBColor(255, 255, 255))
        elif status == "SUSPICIOUS":
            _set_cell_shading(ti_table.rows[row_idx].cells[1], "F1C40F")
            _ct(ti_table.rows[row_idx].cells[1], status, bold=True, size=Pt(9))

    doc.add_paragraph()

    # OTX details
    if otx and not otx.get("error"):
        _add_styled_paragraph(doc, "AlienVault OTX Details", bold=True, size=Pt(11),
                              color=TITLE_COLOR, space_after=Pt(4))
        _add_styled_paragraph(doc, f"Pulse Count: {otx.get('pulse_count', 0)} | "
                              f"Malware Samples: {otx.get('malware_count', 0)} | "
                              f"Malicious URLs: {otx.get('url_count', 0)}", size=Pt(9))

        if otx.get("pulses"):
            pulse_data = []
            for p in otx["pulses"][:10]:
                pulse_data.append((
                    str(otx["pulses"].index(p) + 1),
                    p.get("name", "")[:80],
                    p.get("created", ""),
                    ", ".join(p.get("tags", []))[:60],
                ))
            _make_header_table(doc, "OTX Pulses", ["#", "Pulse Name", "Date", "Tags"],
                               pulse_data, ncols=4)

        if otx.get("malware_samples"):
            doc.add_paragraph()
            _add_styled_paragraph(doc, "Associated Malware:", bold=True, size=Pt(10))
            for m in otx["malware_samples"][:5]:
                _add_styled_paragraph(doc, f"  - {m.get('malware_name', 'Unknown')} "
                                      f"(AV: {m.get('av_name', 'N/A')}, {m.get('date', '')})",
                                      size=Pt(9))

    # AbuseIPDB details
    if abuseipdb and not abuseipdb.get("error"):
        _make_section_table(doc, "AbuseIPDB Details", [
            ("Abuse Confidence Score", f"{abuseipdb.get('abuse_confidence_score', 0)}%"),
            ("Total Reports", str(abuseipdb.get("total_reports", 0))),
            ("Distinct Reporters", str(abuseipdb.get("num_distinct_users", 0))),
            ("Last Reported", abuseipdb.get("last_reported_at", "N/A") or "N/A"),
            ("Country", _country_name(abuseipdb.get("country_code", ""))),
            ("Usage Type", abuseipdb.get("usage_type", "N/A")),
            ("ISP", abuseipdb.get("isp", "N/A")),
        ])
        doc.add_paragraph()

    # VirusTotal details
    if vt and not vt.get("error"):
        mal = vt.get("malicious", 0)
        sus = vt.get("suspicious", 0)
        clean = vt.get("harmless", 0)
        und = vt.get("undetected", 0)
        total_vt = mal + sus + clean + und
        vt_table = _make_section_table(doc, "VirusTotal Details", [
            ("Detection Ratio", f"{mal}/{total_vt} engines"),
            ("Malicious", str(mal)),
            ("Suspicious", str(sus)),
            ("Clean", str(clean)),
            ("Undetected", str(und)),
            ("Reputation Score", str(vt.get("reputation", "N/A"))),
        ])
        # Color detection ratio if malicious
        if mal > 0:
            hex_c, wt = ("E74C3C", True) if mal >= 5 else ("E67E22", True) if mal >= 2 else ("F1C40F", False)
            _set_cell_shading(vt_table.rows[1].cells[1], hex_c)
            tc = RGBColor(255, 255, 255) if wt else RGBColor(0, 0, 0)
            _ct(vt_table.rows[1].cells[1], f"{mal}/{total_vt} engines", bold=True, size=Pt(10), color=tc)
        doc.add_paragraph()

    # Shodan details
    if shodan and not shodan.get("error"):
        shodan_data = [
            ("Open Ports", ", ".join(str(p) for p in shodan.get("ports", [])) or "None"),
            ("OS", shodan.get("os", "N/A") or "N/A"),
            ("Organization", shodan.get("org", "N/A")),
            ("Known CVEs", str(len(shodan.get("vulns", [])))),
        ]
        if shodan.get("vulns"):
            shodan_data.append(("CVE List", ", ".join(shodan["vulns"][:10])))
        _make_section_table(doc, "Shodan Details", shodan_data)
        doc.add_paragraph()

    # ThreatFox details
    if threatfox and not threatfox.get("error"):
        ioc_count = threatfox.get("ioc_count", 0)
        tf_data = [("IOC Associations", str(ioc_count))]
        if ioc_count > 0:
            for ioc in threatfox.get("iocs", [])[:5]:
                tf_data.append((
                    ioc.get("malware", "Unknown"),
                    f"Type: {ioc.get('threat_type', 'N/A')}, Confidence: {ioc.get('confidence', 0)}%"
                ))
        _make_section_table(doc, "ThreatFox Details", tf_data)
        doc.add_paragraph()

    # URLhaus details
    if urlhaus and not urlhaus.get("error"):
        uh_data = [
            ("Listed", "Yes" if urlhaus.get("is_listed") else "No"),
            ("Threat", urlhaus.get("threat", "N/A")),
            ("URL Count", str(urlhaus.get("url_count", 0))),
            ("URLs Online", str(urlhaus.get("urls_online", 0))),
        ]
        _make_section_table(doc, "URLhaus Details", uh_data)
        doc.add_paragraph()

    # ─── Network Reconnaissance ─────────────────────────────────
    if recon:
        # Port scan
        port_scan = recon.get("port_scan", {})
        if port_scan and not port_scan.get("error"):
            open_ports = port_scan.get("open_ports", [])
            if open_ports:
                try:
                    from config import HIGH_RISK_PORTS
                except ImportError:
                    HIGH_RISK_PORTS = {}
                port_data = []
                for p in open_ports:
                    svc = HIGH_RISK_PORTS.get(p, "Unknown")
                    banner = port_scan.get("service_banners", {}).get(p, "")[:60]
                    risk = "HIGH" if p in HIGH_RISK_PORTS else "LOW"
                    port_data.append((str(p), svc, banner or "N/A", risk))

                port_table = _make_header_table(doc, "Port Scan Results",
                                                 ["Port", "Service", "Banner", "Risk"],
                                                 port_data, ncols=4)
                for i, (port, svc, banner, risk) in enumerate(port_data):
                    row_idx = i + 2
                    hex_c, wt = SEV_COLORS.get(risk, ("95A5A6", True))
                    _set_cell_shading(port_table.rows[row_idx].cells[3], hex_c)
                    tc = RGBColor(255, 255, 255) if wt else RGBColor(0, 0, 0)
                    _ct(port_table.rows[row_idx].cells[3], risk, bold=True, size=Pt(9), color=tc)
                doc.add_paragraph()
            else:
                _make_section_table(doc, "Port Scan Results", [
                    ("Open Ports", "0"),
                    ("Scan Time", f"{port_scan.get('scan_time', 0)}s"),
                    ("Result", "No open ports detected"),
                ])
                doc.add_paragraph()

        # Reverse DNS
        rdns = recon.get("reverse_dns", {})
        if rdns:
            hostnames = rdns.get("hostnames", [])
            _make_section_table(doc, "Reverse DNS", [
                ("PTR Records", ", ".join(hostnames) if hostnames else "None"),
            ])
            doc.add_paragraph()

        # HTTP Probe
        http = recon.get("http_probe")
        if http and not http.get("error"):
            http_data = [
                ("HTTP Status", str(http.get("http_status", "N/A"))),
                ("HTTPS Status", str(http.get("https_status", "N/A"))),
                ("Server", http.get("server_header", "N/A") or "N/A"),
                ("TLS Version", http.get("tls_version", "N/A") or "N/A"),
            ]
            sec = http.get("security_headers", {})
            for hdr, val in (sec.items() if sec else []):
                status = "PRESENT" if val != "MISSING" else "MISSING"
                http_data.append((hdr, status))

            _make_section_table(doc, "HTTP/HTTPS Probe", http_data)
            doc.add_paragraph()

    # ─── IOC Table ──────────────────────────────────────────────
    ioc_entries = []
    action = "BLOCK" if risk_cls in ("CRITICAL", "HIGH") else "MONITOR"
    ioc_entries.append((target, target_type.upper(), risk_cls, action, "Primary indicator"))

    if otx and otx.get("malware_samples"):
        for m in otx["malware_samples"][:3]:
            if m.get("hash"):
                ioc_entries.append((m["hash"][:16] + "...", "HASH", "HIGH", "BLOCK",
                                    f"Malware: {m.get('malware_name', 'Unknown')}"))

    if threatfox and threatfox.get("iocs"):
        for ioc in threatfox["iocs"][:3]:
            ioc_entries.append((ioc.get("ioc", "")[:40], ioc.get("ioc_type", ""),
                                "HIGH", "BLOCK", f"Malware: {ioc.get('malware', '')}"))

    if ioc_entries:
        ioc_table = _make_header_table(doc, "IOC Table — Block/Monitor Actions",
                                        ["Indicator", "Type", "Severity", "Action", "Notes"],
                                        ioc_entries, ncols=5)
        for i, (ind, itype, sev, act, notes) in enumerate(ioc_entries):
            row_idx = i + 2
            hex_c, wt = SEV_COLORS.get(sev, ("95A5A6", True))
            _set_cell_shading(ioc_table.rows[row_idx].cells[2], hex_c)
            tc = RGBColor(255, 255, 255) if wt else RGBColor(0, 0, 0)
            _ct(ioc_table.rows[row_idx].cells[2], sev, bold=True, size=Pt(8), color=tc)
            act_color = "E74C3C" if act == "BLOCK" else "F39C12"
            _set_cell_shading(ioc_table.rows[row_idx].cells[3], act_color)
            _ct(ioc_table.rows[row_idx].cells[3], act, bold=True, size=Pt(8),
                color=RGBColor(255, 255, 255))

    doc.add_paragraph()

    # ─── Recommended Actions ────────────────────────────────────
    actions = risk_assessment.get("recommended_actions", [])
    if actions:
        action_data = [(str(i + 1), act) for i, act in enumerate(actions)]
        _make_header_table(doc, "Recommended Actions", ["#", "Action"], action_data, ncols=2)

    doc.add_paragraph()

    # ─── Detection Rules ────────────────────────────────────────
    sigma_rule = f"""title: Traffic to/from {target}
id: {report_id.lower().replace('-', '')}
status: experimental
description: Detects network traffic to/from investigated indicator {target}
references:
    - This threat intelligence report ({report_id})
author: IP/Domain Rep Tool
date: {datetime.now().strftime('%Y/%m/%d')}
logsource:
    category: firewall
detection:
    selection:
        DestinationIp|contains: '{target}' if target_type == 'ip' else ''
        DestinationHostname|contains: '{target}' if target_type == 'domain' else ''
    condition: selection
falsepositives:
    - Legitimate business traffic (review before blocking)
level: {risk_cls.lower()}
tags:
    - attack.command_and_control
    - attack.t1071"""

    splunk_rule = f"""index=* (dest_ip="{target}" OR src_ip="{target}" OR dest="{target}")
| stats count by src_ip, dest_ip, dest, action, app
| sort -count"""

    if target_type == "ip":
        eql_rule = 'network where destination.ip == "' + target + '" or source.ip == "' + target + '"'
        kql_filter = 'source.ip: "' + target + '" OR destination.ip: "' + target + '"'
    else:
        eql_rule = 'network where dns.question.name == "' + target + '" or destination.domain == "' + target + '"'
        kql_filter = 'dns.question.name: "' + target + '" OR destination.domain: "' + target + '"'

    import json as _json
    elastic_obj = {
        "name": "Traffic to/from " + target,
        "description": "Detects connections to investigated indicator " + target,
        "risk_score": risk_assessment.get("score", 0),
        "severity": risk_assessment.get("classification", "low").lower(),
        "type": "eql",
        "query": eql_rule,
        "interval": "5m",
        "from": "now-3600s",
        "enabled": True,
        "tags": ["threat-intel", "soc", "custom"],
        "threat": [{"framework": "MITRE ATT&CK", "tactic": {"id": "TA0011", "name": "Command and Control"}}]
    }

    detection_data = [
        ("Sigma Rule (Firewall)", sigma_rule),
        ("Splunk Query", splunk_rule),
        ("Elastic EQL Rule", eql_rule),
        ("Elastic KQL Filter", kql_filter),
        ("Elastic JSON Import (POST /api/detection_engine/rules)", _json.dumps(elastic_obj, indent=2)),
    ]
    det_table = _make_section_table(doc, "Detection Rules", detection_data)
    # Use monospace-style smaller font for code cells
    for i in range(1, len(det_table.rows)):
        for p in det_table.rows[i].cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.name = "Consolas"

    doc.add_paragraph()

    # ─── Appendix — OSINT Sources ───────────────────────────────
    sources = [
        ("AlienVault OTX", "https://otx.alienvault.com/indicator/ip/" + target),
        ("AbuseIPDB", f"https://www.abuseipdb.com/check/{target}"),
        ("VirusTotal", f"https://www.virustotal.com/gui/ip-address/{target}" if target_type == "ip" else f"https://www.virustotal.com/gui/domain/{target}"),
        ("Shodan", f"https://www.shodan.io/host/{target}" if target_type == "ip" else f"https://www.shodan.io/search?query=hostname:{target}"),
        ("ThreatFox", "https://threatfox.abuse.ch/browse/"),
        ("URLhaus", "https://urlhaus.abuse.ch/browse/"),
        ("IPInfo", f"https://ipinfo.io/{target}"),
    ]
    _make_header_table(doc, "Appendix — OSINT Sources",
                       ["Source", "URL"], sources, ncols=2)

    # ─── Disclaimer Footer ──────────────────────────────────────
    doc.add_paragraph()
    _add_styled_paragraph(doc, "DISCLAIMER", bold=True, size=Pt(9), color=LABEL_COLOR,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    _add_styled_paragraph(doc, (
        "This report was auto-generated by ThreatLens v1.0 from live OSINT queries. "
        "All source data reflects the state at the time of investigation. "
        "Results may differ if queries are re-run at a later time. "
        f"Classification: {classification}."
    ), size=Pt(8), color=RGBColor(0x99, 0x99, 0x99),
       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ─── Save ──────────────────────────────────────────────────
    if not output_path:
        from config import DOWNLOADS_DIR
        safe_target = target.replace(":", "-").replace("/", "-").replace("\\", "-")
        output_path = str(DOWNLOADS_DIR / f"TI_Report_{safe_target}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    try:
        doc.save(output_path)
    except PermissionError:
        output_path = output_path.replace(".docx", "_v2.docx")
        doc.save(output_path)

    return output_path


# ═══════════════════════════════════════════════════════════════════
# TXT Report Generator
# ═══════════════════════════════════════════════════════════════════

def generate_txt_report(
    target: str,
    target_type: str,
    risk_assessment: dict,
    ipinfo: dict = None,
    otx: dict = None,
    abuseipdb: dict = None,
    vt: dict = None,
    shodan: dict = None,
    threatfox: dict = None,
    urlhaus: dict = None,
    recon: dict = None,
    output_path: str = None,
    analyst: str = "Threat Intel Analyst",
    classification: str = "CONFIDENTIAL",
) -> str:
    """Generate a plain-text threat intelligence report (ASCII only)."""
    from hashlib import sha256
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    report_id = f"TI-{datetime.now().strftime('%Y%m%d')}-{target.replace('.', '-').replace(':', '-')}"
    risk = risk_assessment
    W = 72

    lines = []

    def ln(text=""):
        lines.append(text)

    def hdr(title):
        ln("=" * W)
        ln(f"  {title}")
        ln("=" * W)

    def sec(title):
        ln()
        ln(f"--- {title} " + "-" * (W - len(title) - 6))
        ln()

    def kv(key, value):
        ln(f"  {key + ':':<24s} {value}")

    def bar(label, value, max_val=100, width=35):
        filled = int(width * value / max_val) if max_val > 0 else 0
        ln(f"  {label:<20s} {'#' * filled}{'.' * (width - filled)} {value}/{max_val}")

    # ─── Cover ─────────────────────────────────────────────────
    ln()
    ln("=" * W)
    ln("  THREATLENS REPORT")
    ln(f"  IP/Domain Reputation Analysis")
    ln("=" * W)
    ln()
    ln(f"  Investigation ID: {report_id}")
    ln(f"  Target:           {target}")
    ln(f"  Type:             {'IP Address' if target_type == 'ip' else 'Domain'}")
    ln(f"  Risk:             {risk.get('classification', 'N/A')} (Score: {risk.get('score', 0)}/100)")
    ln(f"  Date:             {timestamp}")
    ln(f"  Prepared By:      {analyst}")
    ln(f"  Classification:   {classification}")
    ln()
    ln("=" * W)

    # ─── 1. Executive Summary ──────────────────────────────────
    sec("1. EXECUTIVE SUMMARY")

    signals = risk.get("signals", [])
    not_checked = risk.get("not_checked_sources", [])
    is_good = risk.get("is_known_good", False)

    if is_good:
        ln(f"  STATUS: KNOWN LEGITIMATE — {risk.get('known_good_reason', '')}")
        ln()

    kv("Risk Score", f"{risk.get('score', 0)}/100")
    kv("Classification", risk.get("classification", "N/A"))
    kv("Signals Detected", str(risk.get("signal_count", 0)))
    kv("Known-Good Status", "YES" if is_good else "No")

    if not_checked:
        ln()
        ln(f"  WARNING: The following sources were NOT checked:")
        ln(f"  {', '.join(not_checked)}")
        ln(f"  Assessment may be incomplete.")

    # ─── 2. Source Verification ────────────────────────────────
    sec("2. SOURCE VERIFICATION")
    source_statuses = risk.get("source_statuses", {})
    if source_statuses:
        for src, status in source_statuses.items():
            if status == "SUCCESS":
                marker = "[OK]"
            elif status == "NOT_FOUND":
                marker = "[OK]"  # Successfully queried, just no results
            elif status == "NO_API_KEY":
                marker = "[!!]"
            elif status == "RATE_LIMITED":
                marker = "[!!]"
            elif status == "TIMEOUT":
                marker = "[!!]"
            elif status == "UNAUTHORIZED":
                marker = "[!!]"
            elif status == "FORBIDDEN":
                marker = "[!!]"
            elif status == "SERVER_ERROR":
                marker = "[!!]"
            elif status == "NETWORK_ERROR":
                marker = "[!!]"
            else:
                marker = "[!!]"
            ln(f"  {marker} {src:<16s} {status}")
    else:
        # Fallback: infer from data
        for name, data in [("IPInfo", ipinfo), ("OTX", otx), ("AbuseIPDB", abuseipdb),
                           ("VirusTotal", vt), ("Shodan", shodan),
                           ("ThreatFox", threatfox), ("URLhaus", urlhaus)]:
            if data is None:
                ln(f"  [!!] {name:<16s} NOT RUN")
            elif data.get("error"):
                ln(f"  [!!] {name:<16s} ERROR")
            else:
                ln(f"  [OK] {name:<16s} SUCCESS")

    # Coverage summary
    coverage = risk.get("coverage", {})
    if coverage:
        ln()
        kv("Source Coverage", f"{coverage.get('checked', 0)}/{coverage.get('total', 0)} "
           f"({coverage.get('percentage', 0)}%)")
        kv("Confidence", f"{risk.get('confidence', 0)}%")

    # ─── 3. Indicator Profile ──────────────────────────────────
    sec("3. INDICATOR PROFILE")
    if target_type == "ip" and ipinfo and not ipinfo.get("error"):
        kv("IP Address", target)
        kv("Hostname", ipinfo.get("hostname", "N/A") or "N/A")
        kv("ASN", ipinfo.get("asn", "N/A"))
        kv("ISP/Organization", ipinfo.get("isp", "N/A"))
        kv("Country", _country_name(ipinfo.get("country", "N/A")))
        kv("City/Region", f"{ipinfo.get('city', '')}, {ipinfo.get('region', '')}".strip(", "))
        if ipinfo.get("is_cloud"):
            kv("Cloud Provider", ipinfo.get("cloud_provider", ""))
    elif target_type == "domain" and recon:
        dns = recon.get("dns", {})
        whois_data = recon.get("whois", {})
        kv("Domain", target)
        kv("A Records", ", ".join(dns.get("a_records", [])) or "N/A")
        kv("AAAA Records", ", ".join(dns.get("aaaa_records", [])) or "N/A")
        kv("MX Records", ", ".join(f"{m[0]} (pri {m[1]})" for m in dns.get("mx_records", [])) or "N/A")
        kv("NS Records", ", ".join(dns.get("ns_records", [])) or "N/A")
        kv("Registrar", whois_data.get("registrar", "N/A"))
        kv("Creation Date", whois_data.get("creation_date", "N/A"))
        kv("Expiration Date", whois_data.get("expiration_date", "N/A"))
    else:
        kv("Target", target)

    # ─── 4. Risk Assessment ────────────────────────────────────
    sec("4. RISK ASSESSMENT")

    if signals:
        ln("  Score Breakdown:")
        for sig in signals[:10]:
            bar(sig["source"], sig["weight"])
        ln()

        ln("  Signal Details:")
        ln()
        for i, sig in enumerate(signals, 1):
            tier = sig.get("tier", "?")
            interp = sig.get("interpretation", "")
            ln(f"  {i:2d}. [{sig['severity']}] {sig['source']}: {sig['signal']}")
            ln(f"      Weight: +{sig['weight']}  Tier: {tier}")
            if interp:
                ln(f"      Interpretation: {interp}")
            ln()
    else:
        ln("  No significant threat signals detected.")

    # ─── 5. Mitigating Factors ─────────────────────────────────
    if risk.get("mitigations"):
        sec("5. MITIGATING FACTORS")
        for m in risk["mitigations"]:
            ln(f"  * {m}")

    # ─── 6. Threat Intelligence Findings ───────────────────────
    sec("6. THREAT INTELLIGENCE FINDINGS")

    # OTX
    if otx and not otx.get("error"):
        ln(f"  AlienVault OTX:")
        ln(f"    Status:           MATCH" if otx.get("pulse_count", 0) > 0 else f"    Status:           NO MATCH")
        ln(f"    Pulse Count:      {otx.get('pulse_count', 0)}")
        ln(f"    Malware Samples:  {otx.get('malware_count', 0)}")
        ln(f"    Malicious URLs:   {otx.get('url_count', 0)}")
        if otx.get("pulses"):
            ln(f"    Top Pulses:")
            for p in otx["pulses"][:5]:
                ln(f"      - {p['name']} ({p.get('created', '')})")
                if p.get("tags"):
                    ln(f"        Tags: {', '.join(p['tags'][:5])}")
        if otx.get("malware_samples"):
            ln(f"    Associated Malware:")
            for m in otx["malware_samples"][:5]:
                ln(f"      - {m.get('malware_name', 'Unknown')} (AV: {m.get('av_name', 'N/A')})")
        ln()
    elif otx and otx.get("error"):
        ln(f"  AlienVault OTX: Status: ERROR - {otx.get('error', '')}")
        ln()
    else:
        ln(f"  AlienVault OTX: Status: NOT RUN")
        ln()

    # AbuseIPDB
    if abuseipdb and not abuseipdb.get("error"):
        ln(f"  AbuseIPDB:")
        ln(f"    Status:           MATCH" if abuseipdb.get('abuse_confidence_score', 0) > 0 else f"    Status:           NO MATCH")
        ln(f"    Abuse Confidence: {abuseipdb.get('abuse_confidence_score', 0)}%")
        ln(f"    Total Reports:    {abuseipdb.get('total_reports', 0)}")
        ln(f"    Distinct Users:   {abuseipdb.get('num_distinct_users', 0)}")
        ln(f"    Last Reported:    {abuseipdb.get('last_reported_at', 'N/A') or 'N/A'}")
        ln(f"    Usage Type:       {abuseipdb.get('usage_type', 'N/A')}")
        ln()
    elif abuseipdb and abuseipdb.get("error"):
        ln(f"  AbuseIPDB: Status: ERROR - {abuseipdb.get('error', '')}")
        ln()
    else:
        ln(f"  AbuseIPDB: Status: NOT RUN")
        ln()

    # VirusTotal
    if vt and not vt.get("error"):
        mal = vt.get("malicious", 0)
        sus = vt.get("suspicious", 0)
        total = mal + sus + vt.get("harmless", 0) + vt.get("undetected", 0)
        status_str = "MATCH" if mal > 0 else ("SUSPICIOUS" if sus > 0 else "NO MATCH")
        ln(f"  VirusTotal:")
        ln(f"    Status:           {status_str}")
        ln(f"    Detection Ratio:  {mal}/{total} engines")
        ln(f"    Malicious:        {mal}")
        ln(f"    Suspicious:       {sus}")
        ln(f"    Clean:            {vt.get('harmless', 0)}")
        ln(f"    Undetected:       {vt.get('undetected', 0)}")
        ln(f"    Reputation:       {vt.get('reputation', 'N/A')}")
        ln()
    elif vt and vt.get("error"):
        ln(f"  VirusTotal: Status: ERROR - {vt.get('error', '')}")
        ln()
    else:
        ln(f"  VirusTotal: Status: NOT RUN")
        ln()

    # Shodan
    if shodan and not shodan.get("error"):
        ln(f"  Shodan:")
        ln(f"    Status:           MATCH" if shodan.get('ports') else f"    Status:           NO MATCH")
        ln(f"    Open Ports:       {', '.join(str(p) for p in shodan.get('ports', [])) or 'None'}")
        ln(f"    OS:               {shodan.get('os', 'N/A') or 'N/A'}")
        ln(f"    Organization:     {shodan.get('org', 'N/A')}")
        ln(f"    Known CVEs:       {len(shodan.get('vulns', []))}")
        if shodan.get("vulns"):
            ln(f"    CVE List:         {', '.join(shodan['vulns'][:10])}")
        ln()
    elif shodan and shodan.get("error"):
        ln(f"  Shodan: Status: ERROR - {shodan.get('error', '')}")
        ln()
    else:
        ln(f"  Shodan: Status: NOT RUN")
        ln()

    # ThreatFox
    if threatfox and not threatfox.get("error"):
        ioc_count = threatfox.get("ioc_count", 0)
        if ioc_count > 0:
            ln(f"  ThreatFox:")
            ln(f"    Status:           MATCH")
            ln(f"    IOC Associations: {ioc_count}")
            for ioc in threatfox.get("iocs", [])[:5]:
                ln(f"    - {ioc.get('malware', 'Unknown')} (Type: {ioc.get('threat_type', 'N/A')}, "
                   f"Confidence: {ioc.get('confidence', 0)}%)")
        else:
            ln(f"  ThreatFox:")
            ln(f"    Status:           NO MATCH")
            ln(f"    IOC Associations: 0")
        ln()
    elif threatfox and threatfox.get("error"):
        ln(f"  ThreatFox: Status: ERROR - {threatfox.get('error', '')}")
        ln()
    else:
        ln(f"  ThreatFox: Status: NOT RUN")
        ln()

    # URLhaus
    if urlhaus and not urlhaus.get("error"):
        is_listed = urlhaus.get("is_listed", False)
        if is_listed:
            ln(f"  URLhaus:")
            ln(f"    Status:           MATCH")
            ln(f"    Listed:           YES")
            ln(f"    Threat:           {urlhaus.get('threat', 'N/A')}")
            ln(f"    URLs:             {urlhaus.get('url_count', 0)} (Online: {urlhaus.get('urls_online', 0)})")
        else:
            ln(f"  URLhaus:")
            ln(f"    Status:           NO MATCH")
            ln(f"    Listed:           No")
        ln()
    elif urlhaus and urlhaus.get("error"):
        ln(f"  URLhaus: Status: ERROR - {urlhaus.get('error', '')}")
        ln()
    else:
        ln(f"  URLhaus: Status: NOT RUN")
        ln()

    # ─── 7. Network Reconnaissance ─────────────────────────────
    if recon:
        sec("7. NETWORK RECONNAISSANCE")

        port_scan = recon.get("port_scan", {})
        if port_scan and port_scan.get("open_ports") and not port_scan.get("error"):
            try:
                from config import HIGH_RISK_PORTS
            except ImportError:
                HIGH_RISK_PORTS = {}
            ln(f"  Port Scan Results:")
            ln(f"    Open Ports: {len(port_scan.get('open_ports', []))}")
            ln(f"    Scan Time:  {port_scan.get('scan_time', 0)}s")
            ln()
            ln(f"    {'Port':<8s} {'Service':<16s} {'Risk':<8s} Banner")
            ln(f"    {'-' * 64}")
            for p in port_scan["open_ports"]:
                svc = HIGH_RISK_PORTS.get(p, "Unknown")
                risk_flag = "HIGH" if p in HIGH_RISK_PORTS else "LOW"
                banner = port_scan.get("service_banners", {}).get(p, "")[:40]
                ln(f"    {p:<8d} {svc:<16s} {risk_flag:<8s} {banner}")
            ln()

        rdns = recon.get("reverse_dns", {})
        if rdns:
            ln(f"  Reverse DNS:")
            ln(f"    PTR Records: {', '.join(rdns.get('hostnames', [])) or 'None'}")
            ln()

        http = recon.get("http_probe")
        if http and not http.get("error"):
            ln(f"  HTTP/HTTPS Probe:")
            ln(f"    HTTP Status:  {http.get('http_status', 'N/A')}")
            ln(f"    HTTPS Status: {http.get('https_status', 'N/A')}")
            ln(f"    Server:       {http.get('server_header', 'N/A') or 'N/A'}")
            ln(f"    TLS Version:  {http.get('tls_version', 'N/A') or 'N/A'}")
            sec_hdrs = http.get("security_headers", {})
            if sec_hdrs:
                ln(f"    Security Headers:")
                for h, v in sec_hdrs.items():
                    status = "PRESENT" if v != "MISSING" else "MISSING"
                    ln(f"      {h}: {status}")
            ln()

    # ─── 8. IOC Table ──────────────────────────────────────────
    sec("8. IOC TABLE — BLOCK/MONITOR ACTIONS")
    classification_val = risk.get("classification", "LOW")
    action = "BLOCK" if classification_val in ("CRITICAL", "HIGH") else "MONITOR"
    ln(f"  {'Indicator':<42s} {'Type':<8s} {'Severity':<10s} {'Action':<8s} Notes")
    ln(f"  {'-' * 80}")
    ln(f"  {target:<42s} {target_type.upper():<8s} {classification_val:<10s} {action:<8s} Primary indicator")

    if otx and otx.get("malware_samples"):
        for m in otx["malware_samples"][:3]:
            if m.get("hash"):
                ln(f"  {m['hash'][:40]:<42s} {'HASH':<8s} {'HIGH':<10s} {'BLOCK':<8s} {m.get('malware_name', '')[:25]}")

    if threatfox and threatfox.get("iocs"):
        for ioc in threatfox["iocs"][:3]:
            ln(f"  {ioc.get('ioc', '')[:40]:<42s} {ioc.get('ioc_type', ''):<8s} {'HIGH':<10s} {'BLOCK':<8s} {ioc.get('malware', '')[:25]}")
    ln()

    # ─── 9. Recommended Actions ────────────────────────────────
    sec("9. RECOMMENDED ACTIONS")
    actions = risk.get("recommended_actions", [])
    for i, act in enumerate(actions, 1):
        ln(f"  {i}. {act}")
    ln()

    # ─── 10. Detection Rules ───────────────────────────────────
    sec("10. DETECTION RULES")

    ln("  10.1 SIGMA RULE (Firewall):")
    ln()
    ln(f"  title: Traffic to/from {target}")
    ln(f"  logsource: category: firewall")
    ln(f"  detection:")
    if target_type == "ip":
        ln(f"    selection:")
        ln(f"      DestinationIp|contains: '{target}'")
    else:
        ln(f"    selection:")
        ln(f"      DestinationHostname|contains: '{target}'")
    ln(f"    condition: selection")
    ln(f"  level: {classification_val.lower()}")
    ln()

    ln("  10.2 ELASTIC KQL FILTER:")
    ln()
    if target_type == "ip":
        ln(f"  source.ip: \"{target}\" OR destination.ip: \"{target}\"")
    else:
        ln(f"  dns.question.name: \"{target}\" OR destination.domain: \"{target}\"")
    ln()

    ln("  10.3 ELASTIC EQL RULE:")
    ln()
    ln(f"  // Kibana Security > Rules > Create Custom Rule")
    ln(f"  // Rule type: EQL")
    if target_type == "ip":
        ln(f"  network where destination.ip == \"{target}\" or source.ip == \"{target}\"")
    else:
        ln(f"  network where dns.question.name == \"{target}\"")
        ln(f"    or destination.domain == \"{target}\"")
    ln()

    ln("  10.4 SPLUNK QUERY:")
    ln()
    ln(f"  index=* (dest_ip=\"{target}\" OR src_ip=\"{target}\" OR dest=\"{target}\")")
    ln(f"  | stats count by src_ip, dest_ip, dest, action, app")
    ln(f"  | sort -count")
    ln()

    ln("  10.5 ELASTIC JSON RULE IMPORT (Kibana API):")
    ln()
    ln(f'  POST /api/detection_engine/rules')
    if target_type == "ip":
        eql_q = 'network where destination.ip == "' + target + '" or source.ip == "' + target + '"'
    else:
        eql_q = 'network where dns.question.name == "' + target + '" or destination.domain == "' + target + '"'
    ln(f'  {{')
    ln(f'    "name": "Traffic to/from {target}",')
    ln(f'    "description": "Detects connections to investigated indicator {target}",')
    ln(f'    "risk_score": {risk.get("score", 0)},')
    ln(f'    "severity": "{classification_val.lower()}",')
    ln(f'    "type": "eql",')
    ln(f'    "query": "{eql_q}",')
    ln(f'    "interval": "5m",')
    ln(f'    "from": "now-3600s",')
    ln(f'    "enabled": true,')
    ln(f'    "tags": ["threat-intel", "soc", "custom"],')
    ln(f'    "threat": [{{"framework": "MITRE ATT&CK", "tactic": {{"id": "TA0011", "name": "Command and Control"}}}}]')
    ln(f'  }}')
    ln()

    # --- 11. Appendix ---
    sec("11. APPENDIX — OSINT SOURCES")
    sources = [
        ("AlienVault OTX", f"https://otx.alienvault.com/indicator/ip/{target}"),
        ("AbuseIPDB", f"https://www.abuseipdb.com/check/{target}"),
        ("VirusTotal", f"https://www.virustotal.com/gui/ip-address/{target}" if target_type == "ip" else f"https://www.virustotal.com/gui/domain/{target}"),
        ("Shodan", f"https://www.shodan.io/host/{target}" if target_type == "ip" else f"https://www.shodan.io/search?query=hostname:{target}"),
        ("ThreatFox", "https://threatfox.abuse.ch/browse/"),
        ("URLhaus", "https://urlhaus.abuse.ch/browse/"),
        ("IPInfo", f"https://ipinfo.io/{target}"),
    ]
    for name, url in sources:
        ln(f"  {name:<20s} {url}")
    ln()

    # ─── Footer ────────────────────────────────────────────────
    ln("=" * W)
    ln(f"  Report generated: {timestamp} | Classification: {classification}")
    ln(f"  Investigation ID: {report_id}")
    ln(f"  ThreatLens v1.0")
    ln()
    ln(f"  REPORT INTEGRITY")
    ln(f"  This report was auto-generated from live OSINT queries.")
    ln(f"  All source data reflects the state at the time of investigation.")
    ln(f"  Results may differ if queries are re-run at a later time.")
    ln("=" * W)

    # ─── Save ──────────────────────────────────────────────────
    if not output_path:
        from config import DOWNLOADS_DIR
        safe_target = target.replace(":", "-").replace("/", "-").replace("\\", "-")
        output_path = str(DOWNLOADS_DIR / f"TI_Report_{safe_target}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")

    content = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    return output_path


# ═══════════════════════════════════════════════════════════════════
# Bulk Report Generators
# ═══════════════════════════════════════════════════════════════════

def generate_bulk_txt_report(results_list: list, output_path: str) -> str:
    """Generate a combined TXT report for multiple targets."""
    from collections import Counter
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    W = 72
    lines = []

    def ln(text=""):
        lines.append(text)

    def sec(title):
        ln()
        ln(f"--- {title} " + "-" * (W - len(title) - 6))
        ln()

    # Cover
    ln("=" * W)
    ln("  BULK THREATLENS REPORT")
    ln(f"  {len(results_list)} Target(s) Analyzed")
    ln("=" * W)
    ln(f"  Generated: {timestamp}")
    ln("=" * W)

    # Executive Summary
    sec("EXECUTIVE SUMMARY")
    ln(f"  {'#':<4s} {'Target':<35s} {'Type':<8s} {'Risk':<10s} {'Score':<8s} Status")
    ln(f"  " + "-" * 76)

    for i, r in enumerate(results_list, 1):
        risk = r.get("risk", {})
        cls = risk.get("classification", "N/A")
        score = risk.get("score", 0)
        is_good = risk.get("is_known_good", False)
        status = "KNOWN GOOD" if is_good else cls
        ln(f"  {i:<4d} {r.get('target', '?'):<35s} {r.get('target_type', '?'):<8s} {cls:<10s} {score:<8d} {status}")
    ln()

    cls_counts = Counter(r.get("risk", {}).get("classification", "N/A") for r in results_list)
    good_count = sum(1 for r in results_list if r.get("risk", {}).get("is_known_good"))
    ln(f"  Summary: {len(results_list)} total | "
       f"{cls_counts.get('CRITICAL', 0)} CRITICAL | "
       f"{cls_counts.get('HIGH', 0)} HIGH | "
       f"{cls_counts.get('MEDIUM', 0)} MEDIUM | "
       f"{cls_counts.get('LOW', 0)} LOW | "
       f"{good_count} known-good")

    # Per-Target Details
    for i, r in enumerate(results_list, 1):
        target = r.get("target", "?")
        risk = r.get("risk", {})
        cls = risk.get("classification", "N/A")
        score = risk.get("score", 0)

        sec(f"TARGET {i}: {target}")

        ln(f"  Type: {r.get('target_type', '?').upper()}  |  Risk: {cls} ({score}/100)  |  "
           f"Known-Good: {'YES' if risk.get('is_known_good') else 'No'}")
        ln()

        # Source verification
        source_statuses = risk.get("source_statuses", {})
        if source_statuses:
            ln("  Source Verification:")
            for src, status in source_statuses.items():
                marker = "+" if status == "CHECKED" else "x"
                ln(f"    [{marker}] {src:<16s} {status}")
            ln()

        # IP profile
        ipinfo = r.get("ipinfo")
        if ipinfo and not ipinfo.get("error"):
            ln(f"  ASN: {ipinfo.get('asn', 'N/A')}  |  ISP: {ipinfo.get('isp', 'N/A')}")
            ln(f"  Country: {ipinfo.get('country', 'N/A')}")
            ln()

        # Signals
        signals = risk.get("signals", [])
        if signals:
            ln("  Signals:")
            for sig in signals[:5]:
                tier = sig.get("tier", "?")
                ln(f"    [{sig['severity']}] {sig['source']}: {sig['signal']} (+{sig['weight']}) T{tier}")
                interp = sig.get("interpretation", "")
                if interp:
                    ln(f"      {interp}")
            ln()

        # Mitigations
        if risk.get("mitigations"):
            ln("  Mitigations:")
            for m in risk["mitigations"]:
                ln(f"    * {m}")
            ln()

        not_checked = risk.get("not_checked_sources", [])
        if not_checked:
            ln(f"  NOT CHECKED: {', '.join(not_checked)}")
            ln()

    # Combined IOC Table
    sec("COMBINED IOC TABLE")
    ln(f"  {'Indicator':<38s} {'Type':<8s} {'Risk':<10s} {'Action'}")
    ln(f"  " + "-" * 72)
    for r in results_list:
        target = r.get("target", "?")
        risk = r.get("risk", {})
        cls = risk.get("classification", "LOW")
        action = "BLOCK" if cls in ("CRITICAL", "HIGH") else "MONITOR"
        ln(f"  {target:<38s} {r.get('target_type', '?').upper():<8s} {cls:<10s} {action}")
    ln()

    # Footer
    ln("=" * W)
    ln(f"  Report generated: {timestamp}")
    ln(f"  ThreatLens v1.0 - Bulk Report")
    ln("=" * W)

    content = "\n".join(lines)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    return output_path


def generate_bulk_docx_report(results_list: list, output_path: str) -> str:
    """Generate a combined DOCX report for multiple targets."""
    if not HAS_DOCX:
        raise ImportError("python-docx not installed")

    from collections import Counter

    doc = Document()
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Cover
    _add_styled_paragraph(doc, "Bulk ThreatLens Report",
                          bold=True, size=Pt(24), color=TITLE_COLOR,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(4))
    _add_styled_paragraph(doc, "Confidential \u2014 For Internal Use Only",
                          bold=True, size=Pt(11), color=LABEL_COLOR,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=Pt(12))
    _add_styled_paragraph(doc, f"{len(results_list)} Target(s) Analyzed",
                          size=Pt(14), color=LABEL_COLOR,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_paragraph(doc, f"Generated: {timestamp}", size=Pt(11),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Summary
    cls_counts = Counter(r.get("risk", {}).get("classification", "N/A") for r in results_list)
    summary_data = [
        ("Total Targets", str(len(results_list))),
        ("CRITICAL", str(cls_counts.get("CRITICAL", 0))),
        ("HIGH", str(cls_counts.get("HIGH", 0))),
        ("MEDIUM", str(cls_counts.get("MEDIUM", 0))),
        ("LOW", str(cls_counts.get("LOW", 0))),
        ("Known-Good", str(sum(1 for r in results_list if r.get("risk", {}).get("is_known_good")))),
    ]
    _make_section_table(doc, "Summary", summary_data)

    doc.add_paragraph()

    # Executive summary table
    summary_rows = []
    for i, r in enumerate(results_list):
        risk = r.get("risk", {})
        cls = risk.get("classification", "N/A")
        summary_rows.append((
            str(i + 1),
            r.get("target", "?")[:35],
            r.get("target_type", "?").upper(),
            cls,
            str(risk.get("score", 0)),
        ))

    exec_table = _make_header_table(doc, "Executive Summary",
                                     ["#", "Target", "Type", "Risk", "Score"],
                                     summary_rows, ncols=5)
    for i, r in enumerate(results_list):
        risk = r.get("risk", {})
        cls = risk.get("classification", "N/A")
        hex_c, wt = SEV_COLORS.get(cls, ("95A5A6", True))
        row_idx = i + 2
        _set_cell_shading(exec_table.rows[row_idx].cells[3], hex_c)
        _ct(exec_table.rows[row_idx].cells[3], cls, bold=True, size=Pt(9),
            color=RGBColor(255, 255, 255) if wt else RGBColor(0, 0, 0))

    doc.add_paragraph()

    # Per-target details
    for i, r in enumerate(results_list, 1):
        target = r.get("target", "?")
        risk = r.get("risk", {})
        cls = risk.get("classification", "N/A")

        _add_styled_paragraph(doc, f"Target {i}: {target}",
                              bold=True, size=Pt(14), color=TITLE_COLOR)

        ipinfo = r.get("ipinfo")
        if ipinfo and not ipinfo.get("error"):
            profile = [
                ("Target", target),
                ("Type", r.get("target_type", "?").upper()),
                ("ASN", ipinfo.get("asn", "N/A")),
                ("ISP", ipinfo.get("isp", "N/A")),
                ("Country", _country_name(ipinfo.get("country", "N/A"))),
            ]
            _make_section_table(doc, f"Indicator Profile — {target}", profile)
            doc.add_paragraph()

        signals = risk.get("signals", [])
        if signals:
            _add_styled_paragraph(doc, "Signals:", bold=True, size=Pt(11))
            for sig in signals[:5]:
                tier = sig.get("tier", "?")
                interp = sig.get("interpretation", "")
                _add_styled_paragraph(doc,
                    f"[{sig['severity']}] {sig['source']}: {sig['signal']} (+{sig['weight']}) T{tier}",
                    size=Pt(9))
                if interp:
                    _add_styled_paragraph(doc, f"  {interp}", size=Pt(8),
                                          color=RGBColor(0x99, 0x99, 0x99))

        not_checked = risk.get("not_checked_sources", [])
        if not_checked:
            _add_styled_paragraph(doc, f"NOT CHECKED: {', '.join(not_checked)}",
                                  size=Pt(9), color=RGBColor(0xE6, 0x7E, 0x22))

        actions = risk.get("recommended_actions", [])
        if actions:
            doc.add_paragraph()
            _add_styled_paragraph(doc, "Recommended Actions:", bold=True, size=Pt(10))
            for ai, act in enumerate(actions[:3], 1):
                _add_styled_paragraph(doc, f"  {ai}. {act}", size=Pt(9))

        if i < len(results_list):
            doc.add_paragraph()

    # Combined IOC Table
    ioc_rows = []
    for r in results_list:
        risk = r.get("risk", {})
        cls = risk.get("classification", "LOW")
        action = "BLOCK" if cls in ("CRITICAL", "HIGH") else "MONITOR"
        ioc_rows.append((
            r.get("target", "?")[:35],
            r.get("target_type", "?").upper(),
            cls,
            action,
        ))

    ioc_table = _make_header_table(doc, "Combined IOC Table",
                                    ["Target", "Type", "Risk", "Action"],
                                    ioc_rows, ncols=4)
    for i, r in enumerate(results_list):
        risk = r.get("risk", {})
        cls = risk.get("classification", "LOW")
        action = "BLOCK" if cls in ("CRITICAL", "HIGH") else "MONITOR"
        row_idx = i + 2
        hex_c, wt = SEV_COLORS.get(cls, ("95A5A6", True))
        _set_cell_shading(ioc_table.rows[row_idx].cells[2], hex_c)
        _ct(ioc_table.rows[row_idx].cells[2], cls, bold=True, size=Pt(8),
            color=RGBColor(255, 255, 255) if wt else RGBColor(0, 0, 0))
        act_c = "E74C3C" if action == "BLOCK" else "F39C12"
        _set_cell_shading(ioc_table.rows[row_idx].cells[3], act_c)
        _ct(ioc_table.rows[row_idx].cells[3], action, bold=True, size=Pt(8),
            color=RGBColor(255, 255, 255))

    doc.add_paragraph()
    _add_styled_paragraph(doc, (
        "This report was auto-generated by ThreatLens v1.0. "
        f"Generated: {timestamp}"
    ), size=Pt(8), color=RGBColor(0x99, 0x99, 0x99),
       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if not output_path:
        from config import DOWNLOADS_DIR
        output_path = str(DOWNLOADS_DIR / f"Bulk_Report_{len(results_list)}targets_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    try:
        doc.save(output_path)
    except PermissionError:
        output_path = output_path.replace(".docx", "_v2.docx")
        doc.save(output_path)

    return output_path
